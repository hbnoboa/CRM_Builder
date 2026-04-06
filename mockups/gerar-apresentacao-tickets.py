#!/usr/bin/env python3
"""
Gera documento Word: Sistema de Tickets para Transporte e Logistica
Comparacao direta do modelo atual vs modelo de tickets + por que e melhor.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shading(cell, color):
    s = cell._element.get_or_add_tcPr()
    s.append(s.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):color}))

def p(doc, txt, bold=False, sz=11, color=None, align=None, after=6):
    par = doc.add_paragraph()
    r = par.add_run(txt); r.bold=bold; r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*color)
    if align is not None: par.alignment=align
    par.paragraph_format.space_after=Pt(after)
    return par

def b(doc, txt, prefix=None):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.left_indent = Cm(1.5)
    if prefix:
        r=par.add_run(prefix); r.bold=True; r.font.size=Pt(11)
        par.add_run(txt).font.size=Pt(11)
    else:
        par.add_run(txt).font.size=Pt(11)

def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for pp in c.paragraphs:
            for r in pp.runs: r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(255,255,255)
        shading(c,'1F2937')
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(10)
            if ri%2==1: shading(c,'F3F4F6')
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph()

# ============================================================
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2)
    s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
doc.styles['Normal'].font.name='Calibri'
doc.styles['Normal'].font.size=Pt(11)

# ============================================================
# CAPA
# ============================================================
for _ in range(6): doc.add_paragraph()
p(doc, 'Sistema de Tickets', sz=36, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
p(doc, 'para Transporte e Logistica', sz=22, bold=True, color=(59,130,246), align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
p(doc, 'Antes vs Depois: por que tratar sinistros como tickets', sz=14, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
doc.add_paragraph()
p(doc, 'Exemplo real: Corretora Marisa Dilda (Concordia/SC)', sz=12, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
doc.add_page_break()

# ============================================================
# 1. COMO E HOJE
# ============================================================
doc.add_heading('Como E Hoje', level=1)

p(doc,
    'A Marisa Dilda e uma corretora de seguros de transporte em Concordia/SC '
    'que ja usa o nosso sistema. O sistema funciona, mas gerenciar sinistros '
    'nele hoje e assim:', after=10)

b(doc, '75 campos num formulario so — ', prefix='Formulario gigante: ')
b(doc, 'marca, modelo, placa, ano repetidos pro cavalo + 3 carretas — 12 campos pra veiculos que na maioria das vezes ficam 3/4 vazios', prefix='Veiculos repetidos: ')
b(doc, '14 tipos de documento obrigatorios x 2 campos cada (upload + data) = 28 campos. E nao da pra saber quais faltam sem abrir um por um', prefix='Documentos: ')
b(doc, 'ficam numa aba separada — pra ver o historico completo tem que sair do sinistro', prefix='Follow-ups: ')
b(doc, 'e um dropdown simples (Pendente, Em Andamento, Concluido). Qualquer um muda pra qualquer valor', prefix='Status: ')
b(doc, 'nao tem prioridade, nao tem prazo, nao tem Kanban, nao tem notificacao automatica, nao tem dashboard dedicado', prefix='Gestao de fluxo: ')

p(doc, '', after=6)
p(doc,
    'Funciona? Funciona. Mas a corretora gasta tempo procurando informacao, '
    'perde prazo, nao sabe o status geral dos sinistros e precisa abrir cada '
    'um pra entender o que ta acontecendo.',
    bold=True, after=12)

doc.add_page_break()

# ============================================================
# 2. COMO FICA COM TICKETS
# ============================================================
doc.add_heading('Como Fica Tratando Sinistros Como Tickets', level=1)

p(doc,
    'A ideia: cada sinistro e um ticket. Abre, percorre um fluxo, fecha. '
    'No meio do caminho, todo mundo sabe onde ta.', after=10)

b(doc, '75 → ~30 campos. Veiculos e documentos viram sub-entidades (registros filhos)', prefix='Formulario enxuto: ')
b(doc, 'sub-entidade com 5 campos. Sem limite de veiculos, sem campo vazio', prefix='Veiculos: ')
b(doc, 'sub-entidade com 3 campos + checklist visual dos 14 tipos obrigatorios. Barra "8/14", badge na tabela e Kanban', prefix='Documentos + Checklist: ')
b(doc, 'aparecem na timeline do sinistro junto com eventos automaticos (mudanca de status, doc enviado, alerta de SLA)', prefix='Follow-ups na timeline: ')
b(doc, 'com transicoes controladas (Pendente so vai pra Em Andamento), permissoes (so gerente nega) e cores', prefix='Workflow: ')
b(doc, 'Baixa/Media/Alta/Urgente com timer, alertas em 80% e compliance no dashboard', prefix='Prioridade + SLA: ')
b(doc, 'board visual com 5 colunas, arrasta card = muda status', prefix='Kanban: ')
b(doc, 'SLA violado, doc pendente, status mudou — in-app, email, push', prefix='Notificacoes: ')
b(doc, 'KPIs, graficos por causa/corretor/status, heatmap de SLA, cross-filter', prefix='Dashboards: ')

doc.add_page_break()

# ============================================================
# 3. COMPARACAO DIRETA
# ============================================================
doc.add_heading('Comparacao Direta: Antes vs Depois', level=1)

tbl(doc,
    ['Aspecto', 'Hoje', 'Com Tickets'],
    [
        ['Formulario', '75 campos flat', '~30 campos + 3 sub-entidades'],
        ['Veiculos', '12 campos repetidos (max 4)', 'Sub-entidade: N veiculos sem limite'],
        ['Documentos', '28 campos (14 upload + 14 data)', '3 campos + checklist visual "8/14"'],
        ['Saber quais docs faltam', 'Abrir cada campo', 'Barra de progresso + alerta automatico'],
        ['Follow-ups', 'Aba separada', 'Timeline integrada cronologica'],
        ['Historico completo', 'Nao tem', 'Timeline: follow-ups + eventos automaticos'],
        ['Status', 'Dropdown livre', 'Workflow com transicoes e permissoes'],
        ['Prioridade', 'Nao tem', 'Baixa / Media / Alta / Urgente'],
        ['Prazo', 'Nao tem', 'SLA com timer e alertas'],
        ['Visao geral', 'Contar na lista', 'Kanban com 5 colunas + drag-and-drop'],
        ['Dashboard', 'Nao tem dedicado', 'KPIs, graficos, heatmap, cross-filter'],
        ['Carga de trabalho', 'Perguntar pro corretor', 'Grafico por corretor com ranking'],
        ['Notificacoes', 'Nao tem', 'Multi-canal: in-app, email, push'],
        ['Docs na listagem', 'Nao mostra', 'Badge "8/14" colorido'],
        ['Tags', 'Nao tem', 'Tags coloridas (Carga Total, Judicial...)'],
        ['Mobile', 'Funciona', 'Mesmo + push + timeline'],
    ],
    widths=[3.5, 4, 6.5])

p(doc,
    'Nenhum dado e perdido. Os 75 campos continuam existindo — so reorganizados. '
    'A gestao de fluxo e adicionada por cima.',
    bold=True, after=12)

doc.add_page_break()

# ============================================================
# 4. POR QUE E MELHOR
# ============================================================
doc.add_heading('Por Que E Melhor', level=1)

# --- Checklist resolve o gargalo real ---
doc.add_heading('O checklist resolve o maior gargalo do sinistro', level=2)

p(doc,
    'A SUSEP (Circular 621/2021) define que a seguradora tem 30 dias pra '
    'pagar a indenizacao — mas o relogio so comeca quando TODA a documentacao '
    'e entregue. Se falta um documento, o prazo nao comeca. Se a seguradora '
    'pede complemento, o prazo para.', after=8)

p(doc,
    'Na pratica, a corretora que nao controla quais documentos ja mandou e '
    'quais faltam atrasa o processo inteiro. E como ninguem tem um checklist '
    'de verdade, vira uma troca infinita de email: "manda o BO", "ja mandei", '
    '"nao recebi", "mando de novo".', after=8)

p(doc,
    'Com o checklist visual (barra "8 de 14", alerta automatico pra pendentes, '
    'badge na tabela e no Kanban), a corretora sabe EXATAMENTE o que falta '
    'pra cada sinistro. Manda tudo de uma vez, o relogio de 30 dias comeca, '
    'e o processo anda.',
    bold=True, after=12)

# --- Kanban da visao que ninguem tem ---
doc.add_heading('O Kanban da a visao que ninguem tem hoje', level=2)

p(doc,
    'Hoje, pra saber o status geral dos sinistros, a corretora tem que abrir '
    'a lista, filtrar por status, contar. Pra saber quais estao atrasados, '
    'tem que abrir um por um. Com o Kanban:', after=8)

b(doc, 'Bate o olho e sabe: 12 pendentes, 8 em andamento, 3 concluidos')
b(doc, 'Ve quais cards estao com SLA vermelho (atrasados)')
b(doc, 'Ve quais cards estao com poucos documentos (badge "3/14")')
b(doc, 'Arrasta um card pra outra coluna e o status muda — com workflow controlado')

p(doc, '', after=4)
p(doc,
    'Empresas de logistica que adotaram Kanban tiveram 30% menos atrasos. '
    'Equipes com WIP limits (limite de tickets por coluna) aumentaram '
    'throughput em 40% e reduziram tempo de entrega em 60%.',
    after=12)

# --- SLA controla o que a lei ja exige ---
doc.add_heading('O SLA controla o que a lei ja exige', level=2)

p(doc,
    'A SUSEP ja define um prazo legal de 30 dias. Se a seguradora atrasa, '
    'paga multa de 2% + juros de 1% ao mes. Ou seja, SLA ja existe por lei — '
    'so que ninguem controla de verdade. Nao tem timer, nao tem alerta, nao '
    'tem dashboard mostrando compliance.', after=8)

p(doc,
    'Com SLA no sistema, cada sinistro tem um timer rodando. Alerta em 80%. '
    'Marca como violado quando passa. Dashboard mostra compliance geral. '
    'A corretora consegue provar pra seguradora e pro cliente que ta '
    'cumprindo prazo — ou cobrar quando a seguradora nao cumpre.',
    bold=True, after=12)

# --- Timeline elimina o "abre cada sinistro" ---
doc.add_heading('A timeline elimina o "abre cada sinistro pra ver"', level=2)

p(doc,
    'Hoje, pra saber o que ta acontecendo com um sinistro, a corretora '
    'precisa abrir o registro, ler os campos, ir na aba de follow-ups, '
    'procurar no email se mandou documento, conferir se a seguradora '
    'respondeu. Com a timeline integrada:', after=8)

b(doc, 'Tudo aparece cronologico: follow-ups, docs enviados, mudancas de status, alertas de SLA')
b(doc, 'Nao precisa sair do sinistro pra ver o historico completo')
b(doc, 'Qualquer pessoa da equipe abre o sinistro e entende o contexto em 10 segundos')

p(doc, '', after=4)
p(doc,
    'E o mesmo modelo que Salesforce, HubSpot e Zendesk usam — porque funciona. '
    'Reduce busca de informacao, evita retrabalho e mantem todo mundo alinhado.',
    after=12)

# --- Sub-entidades resolvem o formulario ---
doc.add_heading('As sub-entidades resolvem o formulario gigante', level=2)

p(doc,
    'Formulario com 75 campos e ruim pra todo mundo — pra quem preenche e '
    'pra quem consulta. Pesquisas de UX mostram que cada campo extra reduz '
    'a taxa de preenchimento. Com sub-entidades:', after=8)

b(doc, '12 campos de veiculo viram 5 campos numa sub-entidade (N veiculos sem limite)')
b(doc, '28 campos de documento viram 3 campos + checklist visual')
b(doc, 'Formulario principal cai de 75 pra ~30 campos organizados por categoria')
b(doc, 'Nenhum dado e perdido — e reorganizado de forma mais inteligente')

p(doc, '', after=4)
p(doc,
    'O CRM Builder ja suporta sub-entidades nativamente (parentRecordId). '
    'Nao precisa desenvolver nada novo — so configurar.',
    bold=True, after=12)

# --- Numeros do mercado ---
doc.add_heading('Os numeros mostram a oportunidade', level=2)

tbl(doc,
    ['Dado', 'Valor', 'Fonte'],
    [
        ['Premios de seguro transporte (2024)', 'R$ 6,2 bilhoes', 'CNseg'],
        ['Indenizacoes pagas (2024)', 'R$ 3 bilhoes', 'CNseg'],
        ['Roubos de carga (2024)', '10.478 (27/dia)', 'nstech'],
        ['Prejuizo com roubo (2024)', 'R$ 1,2 bilhao', 'nstech / USP'],
        ['Corretores ativos (SUSEP)', '136.505', 'SUSEP Out/2024'],
        ['Transportadoras (ANTT)', '~285.000', 'RNTRC'],
        ['Prazo SUSEP pra regulacao', '30 dias (da doc completa)', 'Circular 621/2021'],
        ['Maturidade digital do setor', '36 de 100', 'McKinsey'],
        ['Perda por "leakage" em sinistros', '5-11% do custo', 'Hyland / VCA'],
        ['Clientes que trocam apos 1 experiencia ruim', '50%', 'ZipDo'],
        ['Custo manual vs digital por transacao', '3x mais caro', 'CAQH Index'],
    ],
    widths=[5, 3.5, 5.5])

p(doc,
    'Mercado grande, maturidade digital baixa, nenhuma ferramenta combina '
    'CRM de seguros com gestao de tickets. O Quiver (lider, R$ 41 bi em '
    'premios processados) faz carteira e comissao — mas nao faz Kanban, '
    'SLA, timeline nem checklist pra sinistros.',
    bold=True, after=12)

doc.add_page_break()

# ============================================================
# 5. OUTROS SEGMENTOS
# ============================================================
doc.add_heading('Nao E So Pra Corretora', level=1)

p(doc,
    'O mesmo conceito se aplica a qualquer empresa do segmento. O "ticket" '
    'muda de nome, a logica e a mesma:', after=10)

tbl(doc,
    ['', 'Corretora', 'Transportadora', 'Seguradora', 'Gerenciadora'],
    [
        ['O ticket e...', 'Sinistro', 'Ocorrencia de carga', 'Aviso de sinistro', 'Evento monitorado'],
        ['Kanban', 'Aviso → Docs → Analise → Pago', 'Aberta → Investig. → Resolvida', 'Recebido → Regulacao → Pago', 'Detectado → Resolvido'],
        ['SLA', 'Prazo envio docs + SUSEP', 'Prazo resolucao', 'SUSEP 30 dias', 'Minutos (resposta)'],
        ['Checklist', '14 tipos de documento', 'Fotos, BO, relatorio', 'Docs do sinistro', 'Relatorio, GPS'],
        ['Mobile', 'Corretor em campo', 'Motorista abre pelo celular', 'Vistoriador', 'Agente de resposta'],
    ],
    widths=[2.5, 3.2, 3.2, 3, 3])

doc.add_page_break()

# ============================================================
# 6. CONCORRENTES
# ============================================================
doc.add_heading('O Que Existe La Fora', level=1)

tbl(doc,
    ['Ferramenta', 'Faz', 'Nao faz', 'Preco'],
    [
        ['Zendesk', 'Helpdesk generico, chat, tickets', 'Nao entende sinistro, campos limitados', 'USD 55-115/agente/mes'],
        ['Freshdesk', 'Tickets, automacao', 'Sem Kanban custom, sem checklist', 'USD 15-79/agente/mes'],
        ['Jira Service Mgmt', 'ITSM, workflows', 'Focado em TI, complexo, em ingles', 'USD 24-53/agente/mes'],
        ['Quiver (TOTVS)', 'Carteira, cotacao, comissao', 'Sem Kanban, SLA, timeline, checklist', 'Sob consulta'],
        ['Segfy', 'Multicarteira, comissoes', 'Sem tickets, sem fluxo de sinistro', 'R$ 119-150/user/mes'],
        ['Nossa solucao', 'Tickets + sub-entidades + Kanban + SLA + checklist + dashboard + mobile', 'Entende o negocio, tudo integrado', 'Competitivo'],
    ],
    widths=[2.5, 3.5, 4.5, 3.5])

p(doc, '', after=4)
p(doc,
    'Ninguem combina CRM de seguros com gestao de tickets. Os genericos '
    '(Zendesk, Jira) nao entendem sinistro. Os especificos (Quiver, Segfy) '
    'nao fazem gestao de fluxo.',
    bold=True, after=12)

doc.add_page_break()

# ============================================================
# 7. MOCKUP + VIABILIDADE
# ============================================================
doc.add_heading('Mockup e Viabilidade', level=1)

p(doc,
    'Tem um mockup interativo pronto (mockups/ios-risk-mockup.html) com dados '
    'da Marisa Dilda. Da pra navegar, arrastar cards no Kanban, buscar na '
    'tabela, clicar nos graficos, ver o checklist de documentos, timeline, '
    'tudo funcionando.', after=10)

p(doc, 'O que ja ta pronto na plataforma:', bold=True, after=6)
b(doc, 'Sub-entidades, Workflow Status, SLA, Notificacoes, 26 tipos de widget, Tags, 45+ tipos de campo, Automacoes, RBAC, Mobile offline-first, Timeline')

p(doc, '', after=4)
p(doc, 'O que falta:', bold=True, after=6)
b(doc, 'Kanban Board visual (~3-5 dias)')
b(doc, 'Checklist visual de documentos (~2-3 dias)')

p(doc, '', after=6)
p(doc,
    '95% pronto. 1-2 semanas pro resto. Piloto com a Marisa Dilda, '
    'depois escala pra outras corretoras e transportadoras.',
    bold=True, after=10)

# ---------- save ----------
out = os.path.join(os.path.dirname(__file__), 'Apresentacao-Sistema-Tickets.docx')
doc.save(out)
print(f'Documento gerado: {out}')
